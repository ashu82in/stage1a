#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Jun  2 17:13:00 2024

@author: ashutoshgoenka
"""

import pandas as pd
import numpy as np
import streamlit as st
from PIL import Image , ImageOps, ExifTags
import PIL
import os
import zipfile
from zipfile import ZipFile, ZIP_DEFLATED
import pathlib
import shutil
import docx
import docxtpl
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches
import random
from random import randint
from streamlit import session_state
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(layout="wide")

try:
    shutil.rmtree("images_comp_audit")
except:
    pass

try:
    os.mkdir("images_comp_audit")
except:
    pass

def createfile():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    document.save("test.docx")
    
def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width
        

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'bottom', 'end','insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def row_position():
    df_final= df2
    identical_rows_dict = {}
    row_positions  = []
    item_obs = df_final["Element"] + df_final["Observations"]
    # st.write(item_obs.value_counts())
    
    item_obs_list = list(item_obs)
    for i,j in item_obs.value_counts().items():
        identical_rows_dict[i] = j
    #         print(row_positions[-1])
        start_pos = item_obs_list.index(i) + 1
        end_pos = start_pos + j -1
       
        row_positions.append((start_pos, end_pos))
    # st.write(row_positions)
    return row_positions


def allowDocumentBreak(document):
    """Allow table rows to break across pages."""
    tags = document.element.xpath("//w:tr")
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement("w:cantSplit")  # Create arbitrary tag
#         child.set(qn("w:val"), "0")
        tag.append(child)  # Append in the new tag




def updateTable_new():
    createfile()
    updateTable_final()

def updateTable_final():
    global df2
    
    
    doc = docx.Document('test.docx')
    df_final = df2.copy(deep=True)
    df_for_excel = df2.copy(deep=True)
    df_for_excel["No of Images"] = 0
    for idx, rows in df_for_excel.iterrows():
        temp_value = str(rows[9]).split(",")
        temp_value = [t.strip() for t in temp_value]
        df_for_excel.loc[idx, "No of Images"] = len(temp_value)
    df_for_excel["Section"] = 1
    col_idx_temp = ["Sl", "Element", "Observations", "Action Needed", "Location", "Category", "Image No." , "Remarks/ Action By", "Section", "No of Images", "Image Number"]
    df_for_excel = df_for_excel.reindex(columns=col_idx_temp)
    # st.write("in the function")
    # st.write(df_final)
    # st.write(df_for_excel)
    # df_for_excel.to_excel("stage_5_input_table.xlsx")  
    df_final = df_final.drop("Image Number", axis=1)
    t = doc.add_table(df_final.shape[0]+1, df_final.shape[1])
    t.style = 'Table Grid'
    t.allow_autofit = False
    
    for i in range(len(df_final.columns)-1):
        for cell in t.columns[i].cells:
            cell.width = Cm(float(final_col_width[i]))
    
    # for cell in t.columns[0].cells:
    #     cell.width = Cm(1.5)
    # for cell in t.columns[1].cells:
    #     cell.width = Cm(3)
    # for cell in t.columns[2].cells:
    #     cell.width = Cm(3.5)
    # for cell in t.columns[3].cells:
    #     cell.width = Cm(4.5)
    # for cell in t.columns[4].cells:
    #     cell.width = Cm(7.75)
    # for cell in t.columns[5].cells:
    #     cell.width = Cm(2)
    # for cell in t.columns[6].cells:
    #     cell.width = Cm(1.75)
    # for cell in t.columns[7].cells:
    #     cell.width = Cm(2)
    
    # add the header rows.
    for j in range(df_final.shape[-1]):
        t.cell(0,j).text = df_final.columns[j]
    
    # add the rest of the data frame
    for i in range(df_final.shape[0]):
        for j in range(df_final.shape[-1]):
            t.cell(i+1,j).text = str(df_final.values[i,j])
    
    
    row_positions = row_position()
    for row_no in row_positions:
        st_pos = row_no[0]
        end_pos = row_no[1]
        if st_pos != end_pos:
            for i in [0,1,2,3,7]: 
                a = t.cell(st_pos, i)
                temp_content = a.text
    
    #             b = t.cell(2, i)
                c = t.cell(end_pos, i)
            #  # Delete text in cell before merging
            #     delete_paragraph(b.paragraphs[0])
            #     delete_paragraph(c.paragraphs[0])
                c.text = ""
                A = a.merge(c)
                A.text = temp_content
            # add font color to Category
        for col in [4,5,6]:
                for cell in t.columns[col].cells:
                    if cell.text == "Alert":
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(25,140,25)
                    elif cell.text == "Alarm":
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,255)
                    elif cell.text == "Emergency":
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,0,0)
                #     cell.paragraphs[0].paragraph_format.line_spacing = 1.5
                    cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
                    cell.paragraphs[0].paragraph_format.space_after = Cm(0.25)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if col in [5,6]:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        for col in [0,1,2,3,7]:
                for cell in t.columns[col].cells:
                    cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
                    cell.paragraphs[0].paragraph_format.space_after = Cm(0.3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if col in [0,1,7]:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
        for col in [4,5,6]:
                for row in range(st_pos,end_pos+1):
                    temp_cell = t.cell(row, col)
                    if row < end_pos:
                        set_cell_border(
                                temp_cell,
                        #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                                bottom={"sz": 5, "color": "#E6EDF3", "val": "single"},
                #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
                #                 end={"sz": 9, "color": "#000000", "val": "single"},
                            )
                    set_cell_border(
                            temp_cell,
                    #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            #                 bottom={"sz": 9, "color": "#F0F4F8", "val": "single"},
            #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
                            start={"sz": 6, "color": "#000000", "val": "single"},
                            end={"sz": 6, "color": "#000000", "val": "single"},
                        )

            
    # for i in [0,1,2,3,7]:  
    #     a = t.cell(1, i)
    #     temp_content = a.text
        
    #     b = t.cell(2, i)
    #     c = t.cell(3, i)
    # #  # Delete text in cell before merging
    # #     delete_paragraph(b.paragraphs[0])
    # #     delete_paragraph(c.paragraphs[0])
       
    #     A = a.merge(c)
    #     A.text = temp_content
        
    # # add font color to Category
    # for col in [4,5,6]:
    #     for cell in t.columns[col].cells:
    #         if cell.text == "Alert":
    #             cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(25,140,25)
    #     #     cell.paragraphs[0].paragraph_format.line_spacing = 1.5
    #         cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
    #         cell.paragraphs[0].paragraph_format.space_after = Cm(0.25)
    #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #         if col in [5,6]:
    #             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # for col in [0,1,2,3,7]:
    #     for cell in t.columns[col].cells:
    #         cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
    #         cell.paragraphs[0].paragraph_format.space_after = Cm(0.3)
    #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #         if col in [0,1,7]:
    #             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    # for col in [4,5,6]:
    #     for row in range(1,4):
    #         temp_cell = t.cell(row, col)
    #         if row < 3:
    #             set_cell_border(
    #                     temp_cell,
    #             #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
    #                     bottom={"sz": 5, "color": "#E6EDF3", "val": "single"},
    #     #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
    #     #                 end={"sz": 9, "color": "#000000", "val": "single"},
    #                 )
    #         set_cell_border(
    #                 temp_cell,
    #         #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
    # #                 bottom={"sz": 9, "color": "#F0F4F8", "val": "single"},
    # #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
    #                 start={"sz": 6, "color": "#000000", "val": "single"},
    #                 end={"sz": 6, "color": "#000000", "val": "single"},
    #             )
        
        # save the doc
    allowDocumentBreak(doc)
    doc.add_page_break()
    doc.save('./test.docx')   
    


    
    
def updateTable():
    # global up_files
    global final_col_width
    global folder
    global title
    global selection_selected
    # global df_final
    global df
    df_final = df.copy(deep = True)
    
    document = Document("Audit_Word.docx")
    section_selected ="Capstone"
    document.add_heading(section_selected, 1)

    # document.add_paragraph(section_selected)
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = document.add_table(df_final.shape[0]+1, df_final.shape[1])
    t.style = 'Table Grid'
    t.allow_autofit = False
    t.columns[1].width = Cm(7.5)
    st.write(df_final.shape)
    # add the header rows.
    for j in range(df_final.shape[-1]):
        t.cell(0,j).text = df_final.columns[j]
    
    # add the rest of the data frame
    for i in range(df_final.shape[0]):
        for j in range(df_final.shape[-1]):
            t.cell(i+1,j).text = str(df_final.values[i,j])
    
    set_column_width(t.columns[1], docx.shared.Cm(7.5))
    set_column_width(t.columns[2], docx.shared.Cm(5.5))
    set_column_width(t.columns[3], docx.shared.Cm(2))
    # save the doc
    # document.add_paragraph('')
    
    
    
    document.add_heading(section_selected + " - Images", 2)
    
    _, _, files = next(os.walk(folder))
    file_count = len(files)
    st.write(file_count)
    no_of_rows = int(((file_count-1)//3+1)*2)
    
    
    table = document.add_table(rows = no_of_rows , cols = 3)
    # st.write("Table Rows " + str(table.rows.))
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Item'     
    # hdr_cells[1].text = 'quantity'
    document.save("Audit_Word.docx")
    counter = 0
    counter_cols = 0
    title = 1
    for file in folder.iterdir():
        name = os.path.splitext(file.name)[0]
        img_no = int(name.split(" ")[1])
        adj_img_no = img_no  - int(title)
        # st.write(img_no, adj_img_no)
        row_no = (adj_img_no//3) *2
        col_no = int(adj_img_no - (row_no*3/2))
        # if(row_no>0 and col_no==0):
        #     table.add_row()
        #     table.add_row()
        # st.write(img_no, adj_img_no, row_no, col_no)
        # cell = table.rows[counter].cells[counter_cols]
        cell = table.rows[row_no].cells[col_no]
        cell._element.clear_content()
        st.write(file.name + "   a")
        picture = cell.add_paragraph().add_run().add_picture('images_comp_audit/'+file.name, width=Inches(2.6))
        cell = table.rows[row_no+1].cells[col_no]
        # cell = table.rows[counter+1].cells[counter_cols]
        # st.write(row_no, col_no)
        cell.text = name
        if col_no<2:
            counter_cols = counter_cols + 1
        else:
            # table.add_row()
            counter_cols = 0
            counter = counter+2
    document.add_page_break()
    document.save("Audit_Word.docx")

# def resize_image(img, width, height):

#     # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
#     box=((img.width-width)/2,((img.height-height)/2),(img.width+width)/2,((img.height+height)/2))
#     im_resized = im.crop(box)
#     return im_resized

st.title("Stage 1a - Upload Observation - No Sorting")
obs_file = st.file_uploader("Upload Observation Excel File", type=['csv','xlsx'],accept_multiple_files=False,key="fileUploader", on_change=createfile)
if obs_file is not None:
    # uplaod remedy file
    df_rem = pd.read_excel("remedy_excel.xlsx")
    remedy_dict = {}
    for idx, val in df_rem.iterrows():
    #     print(val)
        temp_key = val["Observations"] + val["Severity"]
        remedy_dict[temp_key] = (val["Remedy"], val["Category"])
        
    
    df = pd.read_excel(obs_file)
    df = df.dropna(thresh=5)
    df["obs_element"] = df["Element"] + "_"+df["Observations"] + df["Severity"]
    df["Location_Final"] =  df["Location"] + " at "+ df["Level"]
    df["Action Needed"] = ""
    df["Category"] = ""
    # df = df.sort_values("obs_element")
    df["Image Number"] =  df["Image Number"].astype(str)
    df["Image No."] = ""
    # df["Image List"] = ""
    df["Start Image Number"] = 0
    df["End Image Number"] = 0
    image_col = df.columns.get_loc("Image Number")
    # st.write(image_col)
    st.title("Input Data")
    st.write(df)
    start_val = 1
    end_val = 1
    img_master_list = []
    img_new_old_dict = {}
    location_dict = {}
    counter_temp =1
    missing_temp_data = []
    for idx, row in df.iterrows():
        temp_key = row["Observations"] + row["Severity"]
        try:
            df.loc[idx,"Action Needed"] = remedy_dict[temp_key][0]
            df.loc[idx,"Category"] = remedy_dict[temp_key][1]
        except:
            missing_temp_data.append(temp_key)
        
        df.loc[idx,"Sl"] = counter_temp
        counter_temp += 1
        if row["Severity"] !="" :
            df.loc[idx,"Observations"] = row["Observations"] +" - " + row["Severity"] 
        # st.write(idx)
        # st.write(row[image_col])
        # st.write(df.loc[1,"Observations"])
        temp_value = row[image_col].split(",")
        temp_value = [t.strip() for t in temp_value]
        # st.write(temp_value)
        img_master_list = img_master_list + temp_value
        no_of_img = int(len(temp_value))
        # df.loc[idx, "Image List"] = temp_value
        df.loc[idx, "No of Images"] = no_of_img
        df.loc[idx, "Start Image Number"]  = start_val
        df.loc[idx, "End Image Number"]  = start_val+no_of_img -1
        if no_of_img>1:
            df.loc[idx, "Image No."] = "00" + str(start_val) +" - 00" + str(start_val+no_of_img -1)
        else:
            df.loc[idx, "Image No."] = "00" + str(start_val)
        for ctr in range(no_of_img):
            img_new_old_dict[temp_value[ctr]] = start_val + ctr
            location_dict[start_val + ctr] = df.loc[idx, "Location_Final"] 
                
        start_val = start_val+no_of_img
        end_val = start_val
    # st.write(list(set(missing_temp_data)))
    df2 = df.copy(deep=True)
    
    df2 = df2.drop(["Location", "Level", "Severity", "obs_element", "Start Image Number", "End Image Number", "No of Images"], axis=1)
    
    try:
        df2 = df2.drop(["Section"], axis=1)
    except:
        pass
    df2.rename(columns = {'Location_Final':'Location'}, inplace = True)
    col_idx = ["Sl", "Element", "Observations", "Action Needed", "Location", "Category", "Image No." , "Remarks/ Action By", "Image Number"]
    df2 = df2.reindex(columns=col_idx)
    # st.write(df)
    st.title("Do you want to Modify the Location")
    # df2_mod = df2.iloc[:, :-1]
    df2 = st.data_editor(
            df2,
            column_config={
                "Location": st.column_config.Column(
                    "Location",
                    help="The order in which the table will be added to word",
                    # min_value=1,
                    # max_value=len(location_list),
                    # step=1,
                    # format="%d"
                    required=True,
                        )
                    },
                    hide_index=True,
                    )
    
    st.title("Please Verify the Final Output Data (This will be exported to Word File")
    st.write(df2.iloc[:, :-1])
    # st.write(df2_mod)
    
    
    st.write("Table Column Width in cm")
    no_of_cols = len(df2.columns)
    default_co_width = [1.5, 3.0, 3.5, 4.5, 7.75, 2, 1.75,2]
    col_list = st.columns(no_of_cols-1)
    final_col_width= [i for i in default_co_width]
    total_width = sum([float(i) for i in final_col_width])
    for i in range(no_of_cols-1):
        with col_list[i]:
            final_col_width[i] = st.text_input(df2.columns[i],
                          default_co_width[i],
                          key="col_data_"+str(i),
    )
    if final_col_width[0]:
        total_width = sum([float(i) for i in final_col_width])
    st.write("Total Column Width: "+ str(total_width))
    # st.write(final_col_width)
    # st.write(img_master_list)
    # st.write(img_new_old_dict)
    # st.write(location_dict)
    # up_files = st.file_uploader("Upload Image Files", type = ["png", "jpeg", "jpg"] ,accept_multiple_files=True)
    
    # if up_files is not None:
    #     file_name_list = [file.name for file in up_files]
    #     file_not_found = []
    #     file_found = []
    #     image_file_dict_final = {}
    #     file_image_dict= {}
    #     for img_name in img_master_list:
    #         temp_found = [s for s in file_name_list if img_name in s]
    #         if len(temp_found)>0:
    #             file_name_list = file_name_list + temp_found
    #             image_file_dict_final[img_new_old_dict[img_name]] = temp_found[0]
    #             file_image_dict[temp_found[0]] = img_new_old_dict[img_name]
    #         else:
    #             file_not_found.append(img_name)
        
    #     st.write(image_file_dict_final)
    #     st.write(file_not_found)
    #     st.write(file_found)
    #     st.write(file_image_dict)
    #     image_number_file_dict = {}
        
    #     # directory_for_input = "images_comp_audit"
    #     # folder = pathlib.Path(directory_for_input)
    #     for temp_file in up_files:
    #         try:
    #             st.write(temp_file.name)
    #             # st.write(file_image_dict[temp_file.name])
    #             img_no = file_image_dict[temp_file.name] 
    #             st.write(img_no)
    #             ext= temp_file.name.split(".")[-1]
    #             im = Image.open(temp_file)
    #             # This will keep the orientation of the image as per the original file
    #             for orientation in ExifTags.TAGS.keys():
    #                 if ExifTags.TAGS[orientation]=='Orientation':
    #                     break
    #             exif=dict(im._getexif().items())
    #             if exif[orientation] == 3:
    #                 im=im.rotate(180, expand=True)
    #             elif exif[orientation] == 6:
    #                 im=im.rotate(270, expand=True)
    #             elif exif[orientation] == 8:
    #                 im=im.rotate(90, expand=True)
                
    #             # im.save("images_comp_audit/Image "+str(img_no)+"."+ext, exif=im.info.get("exif"))
    #             im.save("images_comp_audit/Image "+str(img_no)+"."+ext)
                
                
    #         except:
    #             pass
            
        
#         try:
#     image=Image.open(filepath)

#     for orientation in ExifTags.TAGS.keys():
#         if ExifTags.TAGS[orientation]=='Orientation':
#             break
    
#     exif = image._getexif()

#     if exif[orientation] == 3:
#         image=image.rotate(180, expand=True)
#     elif exif[orientation] == 6:
#         image=image.rotate(270, expand=True)
#     elif exif[orientation] == 8:
#         image=image.rotate(90, expand=True)

#     image.save(filepath)
#     image.close()
# except (AttributeError, KeyError, IndexError):
#     # cases: image don't have getexif
#     pass
        
            
        # original_image_size = {}
        # image_size_dict = {}
        # new_width_dict = {}
        # new_height_dict = {}
        
        # directory_for_input = "images_comp_audit"
        # folder = pathlib.Path(directory_for_input)
        
        
        # for file in folder.iterdir():
        #     temp_no = file.name.split(".")[0]
        #     temp_no_1 = temp_no.split(" ")[1]
        #     image_number_file_dict[temp_no_1]  = file
        
        
        # for k,file in image_number_file_dict.items():
            
        #     extensions = ["jpg", "jpeg", "png", "gif", "webp"]
        #     im = Image.open(file)
        #     im = ImageOps.exif_transpose(im)
        #     ext = file.name.split(".")[-1]
            
            
            
        #     # Displaying Image
        #     im_width, im_height = im.size 
        #     original_image_size[file.name] = [im_width, im_height]
        #     try:
        #         b = image_size_dict[file.name]
        #     except:
        #         image_size_dict[file.name] = [im_width, im_height]
                
            
                
            
        #     st.write(im_width, im_height)
        #     size_to_scale = min(im_width,im_height)
        #     st.write(size_to_scale)
        #     # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
        #     box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
        #     im_resized = im.crop(box)
            
        #     col1, col2, col3  = st.columns(3)
        #     with col1:
        #         st.image(im, width=350)
        #         # oi_width  = st.number_input("width", value = im_width)
        #         # oi_height = st.number_input("height", value = im_height)
        #         st.write(im_width, im_height)
                
            
            
        #     with col2:
        #         try:
        #             im_resized = resize_image(im, new_width_dict[file.name], new_height_dict[file.name])
        #         except:
        #             im_resized = im_resized
                    
                
        #         col2_img = st.image(im_resized, width=350)
        #         st.write(im_resized.size)
        #         new_width_dict[file.name]  = st.number_input("new width", value = im_resized.size[0], key="nw_"+file.name)
        #         new_height_dict[file.name] = st.number_input("new height", value = im_resized.size[1], key="nh_"+file.name)
        #         # st.write(im_width, im_height)
        #         st.write(im_resized.size)
            
            
        #     with col3:
        #         try:
        #             im_resized_final = resize_image(im, new_width_dict[file.name], new_height_dict[file.name])
        #         except:
        #             im_resized_final = im_resized
                
                
        #         col3_img = st.image(im_resized_final, width=350)
        #         st.write(im_resized_final.size)
                    
        #     st.write(file.name)
        #     # st.write(location_dict[1])
        #     st.write(str(k))
        #     st.write(location_dict[int(k)])
            
         
        #     im_resized_final.save("images_comp_audit/"+file.name)
            
            
            
        
        # zip_path = "images_compressed_audit.zip"
        # directory_to_zip = "images_comp_audit"
        # folder = pathlib.Path(directory_to_zip)
        # with ZipFile(zip_path, 'w', ZIP_DEFLATED) as zip:
        #     for file in folder.iterdir():
        #         zip.write(file, arcname=file.name)
                
        # with open("images_compressed_audit.zip", "rb") as fp:
        #     btn = st.download_button(
        #         label="Download ZIP",
        #         data=fp,
        #         file_name="images_compressed_audit.zip",
        #         mime="application/zip"
        #     )
        
    if obs_file is not None:
            # df_final = df2.copy(deep=True)
            df_for_excel = df2.copy(deep=True)
            df_for_excel["No of Images"] = 0
            # st.write(df_for_excel)
            for idx, rows in df_for_excel.iterrows():
                temp_value = str(rows[8]).split(",")
                temp_value = [t.strip() for t in temp_value]
                df_for_excel.loc[idx, "No of Images"] = len(temp_value)
            df_for_excel["Section"] = 1
            col_idx_temp = ["Sl", "Element", "Observations", "Action Needed", "Location", "Category", "Image No." , "Remarks/ Action By", "Section", "No of Images", "Image Number"]
            df_for_excel = df_for_excel.reindex(columns=col_idx_temp)
            # st.write("in the function")
            # st.write(df_final)
            # st.write(df_for_excel)
            df_for_excel.to_excel("stage_5_input_table.xlsx")  
            try:
                with open("test.docx", "rb") as fp:
                
                    btn_1 = st.button(
                            label="Create New Word File and Update Data",
                            on_click=updateTable_new       
                        )
            except:
                pass
            
        
        
        
            try:
                with open("test.docx", "rb") as fp:
                
                    btn_1 = st.button(
                            label="Update an Existing Word File",
                            on_click=updateTable_final       
                        )
            except:
                pass
                # st.write(btn_1)
                
                # if btn_1:
                #     st.write("Running Update Function")
                #     updateTable(up_files)
            
            try:
                with open("test.docx", "rb") as fp:
                
                    btn_1 = st.download_button(
                            label="Download Word File",
                            data=fp,
                            file_name="test",
                            mime="docx"
                            )
            except:
                pass
            
            
            try:
                with open("stage_5_input_table.xlsx", "rb") as template_file:
                    template_byte = template_file.read()
                    btn_1 = st.download_button(
                            label="Download Stage 5 Input Excel File",
                            data=template_byte,
                            file_name="stage_5_input_table.xlsx",
                            mime='application/octet-stream'
                            )
            except:
                pass


    # os.remove(zip_path)
            
            
            
            
        
        
# st.write(up_files)
    
    